from docx import Document

# Load an existing Word document
doc = Document('test.docx')

# Iterate through paragraphs and print text
for para in doc.paragraphs:
    print("XXXXXXXXXX")
    print(para.text)

core_props = doc.core_properties
print (core_props) 

print (core_props.title)


import json
from docx import Document

def docx_to_json(file_path):
    doc = Document(file_path)
    data = {
        'paragraphs': [para.text for para in doc.paragraphs]
    }
    return json.dumps(data, indent=4)

# Usage
file_path = 'test.docx'
json_data = docx_to_json(file_path)
print(json_data)
